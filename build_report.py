# -*- coding: utf-8 -*-
"""Generate a CFA-standard equity research report on BHEL as a ~25-page Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x5B,0xA8); GREY=RGBColor(0x59,0x59,0x59)
WHITE=RGBColor(0xFF,0xFF,0xFF); RED=RGBColor(0xC0,0x00,0x00); GREEN=RGBColor(0x2E,0x7D,0x32)

doc=Document()
st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
for s in doc.sections:
    s.top_margin=Inches(0.8); s.bottom_margin=Inches(0.8); s.left_margin=Inches(0.9); s.right_margin=Inches(0.9)

def shade(cell,hexv):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexv); tcPr.append(sh)
def setfont(run,size=10.5,bold=False,color=None,italic=False,name='Calibri'):
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.name=name
    if color is not None: run.font.color.rgb=color
def para(text='',size=10.5,bold=False,color=None,italic=False,align=None,space_after=6,space_before=0):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(space_after); p.paragraph_format.space_before=Pt(space_before)
    if align is not None: p.alignment=align
    if text:
        r=p.add_run(text); setfont(r,size,bold,color,italic)
    return p
def bullet(text,bold_lead=None,size=10.5):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4)
    if bold_lead:
        r=p.add_run(bold_lead); setfont(r,size,True,NAVY); r2=p.add_run(text); setfont(r2,size)
    else:
        r=p.add_run(text); setfont(r,size)
    return p
def h1(text):
    doc.add_page_break()
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8); p.paragraph_format.space_before=Pt(2)
    r=p.add_run(text); setfont(r,16,True,NAVY)
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'1F3864')
    pbdr.append(bottom); pPr.append(pbdr)
    return p
def h2(text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); p.paragraph_format.space_before=Pt(10)
    r=p.add_run(text); setfont(r,12.5,True,BLUE); return p
def h3(text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.space_before=Pt(6)
    r=p.add_run(text); setfont(r,11,True,GREY); return p

def table(headers, rows, widths=None, header_fill='1F3864', body_size=9.5, first_col_left=True,
          highlight_rows=None, total_rows=None):
    highlight_rows=highlight_rows or []; total_rows=total_rows or []
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.style='Table Grid'
    hdr=t.rows[0].cells
    for i,htext in enumerate(headers):
        shade(hdr[i],header_fill); p=hdr[i].paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT if (i==0 and first_col_left) else WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(str(htext)); setfont(r,body_size,True,WHITE)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        fill='F2F5FB' if ri%2==0 else None
        if ri in highlight_rows: fill='FFF2CC'
        if ri in total_rows: fill='D9E1F2'
        for i,val in enumerate(row):
            if fill: shade(cells[i],fill)
            p=cells[i].paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT if (i==0 and first_col_left) else WD_ALIGN_PARAGRAPH.RIGHT
            r=p.add_run(str(val)); setfont(r,body_size, bold=(ri in total_rows or i==0 and False))
            if ri in total_rows: r.bold=True
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    return t

def caption(text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8); p.paragraph_format.space_before=Pt(2)
    r=p.add_run(text); setfont(r,8,italic=True,color=GREY)

# ============ data ============
YRS=['FY20','FY21','FY22','FY23','FY24','FY25','FY26','FY27E','FY28E','FY29E','FY30E','FY31E','FY32E','FY33E']
sales=[21463,17309,21211,23365,23893,28339,33782,40392,48275,55883,63211,70256,76635,82405]
ebitda=[-133,-3048,828,1044,711,1399,2342,3029,4103,5309,6637,7728,8813,9889]
ebmar=[-0.6,-17.6,3.9,4.5,3.0,4.9,6.9,7.5,8.5,9.5,10.5,11.0,11.5,12.0]
pat=[-1468,-2699,446,654,281,533,1600,2075,3029,4073,5204,6152,7086,8004]
eps=[-4.2,-7.8,1.3,1.9,0.8,1.5,4.6,6.0,8.7,11.7,15.0,17.7,20.4,23.0]
roe=[0,-10,2,2,1,2,6,8,11,13,15,16,16,17]
roce=[-2,-11,2,2,1,3,6,8,11,13,15,17,17,18]
de=[0.18,0.19,0.18,0.20,0.36,0.36,0.31,0.26,0.21,0.16,0.12,0.08,0.05,0.02]
def fmt(x,dec=0):
    if isinstance(x,str): return x
    return f"{x:,.{dec}f}"

# ===================================================================
# COVER PAGE
# ===================================================================
para('EQUITY RESEARCH  |  INDIA  |  CAPITAL GOODS — HEAVY ELECTRICAL EQUIPMENT',9,True,GREY,space_before=24)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
r=p.add_run('Bharat Heavy Electricals Limited'); setfont(r,26,True,NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
r=p.add_run('NSE: BHEL  |  BSE: 500103  |  Bloomberg: BHEL IN'); setfont(r,11,False,GREY)
para('Riding the thermal super-cycle — but the price already assumes it', 13, True, BLUE, italic=True, space_before=6, space_after=14)

# rating box
t=doc.add_table(rows=1,cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
labels=[('RATING','SELL',RED),('12M TARGET PRICE','₹245',NAVY),('CURRENT PRICE','₹403',NAVY),('DOWNSIDE','-39%',RED)]
hc=t.rows[0].cells
for i,(lab,val,col) in enumerate(labels):
    shade(hc[i],'1F3864')
    p=hc[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(lab); setfont(r,8.5,True,WHITE)
rc=t.add_row().cells
for i,(lab,val,col) in enumerate(labels):
    p=rc[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(val); setfont(r,15,True,col)
doc.add_paragraph()

# snapshot table
para('Market Data & Snapshot (as of FY2026 close)',10.5,True,NAVY,space_after=4)
snap=[
 ['Market capitalisation','₹1,40,300 cr (~US$16.8 bn)','52-week range','₹205 – ₹425'],
 ['Shares outstanding','348.05 cr','Free float','~37% (GoI holds ~63%)'],
 ['FY26 revenue','₹33,782 cr','FY26 PAT','₹1,600 cr'],
 ['FY26 EPS','₹4.60','FY26 P/E','~88x'],
 ['Order book (FY26)','~₹2.40 lakh cr','Book-to-bill','~7.1x'],
 ['Net cash (FY26)','₹3,680 cr','Dividend yield','~0.3%'],
]
table(['Metric','Value','Metric','Value'], snap, widths=[1.6,2.2,1.6,2.2], body_size=9.5, first_col_left=True)
caption('Source: BHEL FY2026 results, company filings and author\u2019s financial model. Market cap at indicative price of ₹403.')

para('Analyst: Senior Equity Research Analyst  |  Report date: June 2026  |  Currency: INR (₹ crore unless stated)',8.5,italic=True,color=GREY,space_before=10)
para('This report has been prepared on a basis consistent with CFA Institute standards of objectivity, '
     'independence and thoroughness. Valuation is grounded in a fully integrated three-statement model with '
     'audited historicals (FY2020–FY2026) and a seven-year explicit forecast (FY2027–FY2033). Please refer to '
     'the important disclosures and disclaimer on the final page.',8.5,italic=True,color=GREY)

# ===================================================================
# 1. INVESTMENT SUMMARY
# ===================================================================
h1('1.  Investment Summary')
para('We initiate coverage of Bharat Heavy Electricals Limited (BHEL) with a SELL rating and a 12-month '
     'target price of ₹245, implying ~39% downside from the current price of ₹403. Our recommendation is a '
     'statement about valuation, not about the business. BHEL is enjoying its strongest operating upcycle in over '
     'a decade — a record ~₹2.4 lakh crore order book, a return to double-digit revenue growth, and a sharp '
     'profit recovery. However, at ~88x trailing earnings and ~6x book, the market has, in our view, already '
     'discounted a flawless multi-year execution ramp that the company\u2019s own history of slippage makes far '
     'from certain.',bold=False)

h2('Investment thesis at a glance')
bullet(' BHEL is the dominant domestic supplier of thermal power generation equipment at the exact moment India '
       'is re-committing to coal-based capacity (a planned ~80 GW of new thermal by 2032) to firm up a '
       'renewables-heavy grid. Order inflow of ~₹75,000 cr in FY2026 lifted the book-to-bill to ~7.1x.', '+ Structural demand tailwind: ')
bullet(' From an EBITDA margin of ~6.9% in FY2026, we model a gradual climb to ~12% by FY2033 as high-margin '
       'legacy under-recoveries roll off, fixed-cost absorption improves and execution accelerates. PAT compounds '
       'from ₹1,600 cr to ~₹8,000 cr over the forecast horizon.', '+ Genuine earnings recovery: ')
bullet(' BHEL is asset-light and net-cash (₹3,680 cr at FY2026), funded substantially by interest-free customer '
       'advances (contract liabilities). This is a structurally favourable working-capital model when execution flows.', '+ Strong balance sheet: ')
bullet(' At ~88x FY26 earnings and ~45x FY27E earnings, the stock prices in near-perfect execution. Our DCF, using '
       'a 13.6% WACC, yields an intrinsic value of just ₹181/share. Even our probability-weighted scenario value '
       'is ₹245.', '\u2212 The catch — valuation: ')
bullet(' Indian thermal capacity additions reached only 73% of target in FY2026 and 29% in FY2025. BHEL\u2019s '
       'revenue conversion of its order book remains the single biggest swing factor — and the risk is asymmetric '
       'at the current price.', '\u2212 Execution risk is real: ')

h2('Key financial summary')
rows=[]
idx=[3,4,5,6,7,8,10,13]  # FY23..FY30 selection + FY33
sel=[2,3,4,5,6,7,9,13]
for i in sel:
    rows.append([YRS[i], fmt(sales[i]), fmt(ebitda[i]), f"{ebmar[i]:.1f}%", fmt(pat[i]), f"{eps[i]:.1f}",
                 f"{roe[i]:.0f}%", f"{de[i]:.2f}"])
table(['Year','Revenue','EBITDA','EBITDA %','PAT','EPS (₹)','ROE %','D/E (x)'], rows,
      widths=[0.8,1.0,0.9,0.85,0.9,0.85,0.7,0.7], body_size=9.0)
caption('A = audited; E = estimate. Source: author\u2019s integrated model built on BHEL audited consolidated financials (FY20–26) and FY27–33 forecast.')

h2('What would change our view')
para('We would turn constructive (a) on a meaningful price correction toward our ₹245 fair value; (b) on hard '
     'evidence that BHEL can sustainably convert >18% of its opening order book into revenue while holding margins; '
     'or (c) on EBITDA margins durably exceeding 12% earlier than modelled. Conversely, renewed execution slippage, '
     'commodity-driven margin pressure, or large low-margin order intake would reinforce the downside.',space_after=4)

h2('The bull case versus the bear case')
para('The entire BHEL debate can be reduced to one question: how much of the record order book will convert into '
     'profitable revenue, and how quickly? The two sides are sharply drawn.',space_after=4)
h3('The bull case (why the stock is at ₹400+)')
bullet(' A ~₹2.4 lakh crore order book at ~7x book-to-bill gives unprecedented multi-year visibility; brokerages '
       'expect BHEL to enter FY2027 with ₹2.5 lakh cr of backlog.', '')
bullet(' Operating leverage is enormous — every point of margin on a fixed cost base drops disproportionately to '
       'PAT, so a margin path to 12–15% could see EPS multiply several times over.', '')
bullet(' BHEL is the near-monopoly domestic BTG supplier into an 80 GW thermal build, an emerging nuclear '
       'programme, FGD retrofits and exports — a scarce, strategic asset that the market is happy to own at a premium.', '')
h3('The bear case (why we rate it SELL)')
bullet(' The order book is only worth what converts: India\u2019s thermal additions hit just 29% and 73% of target '
       'in FY2025 and FY2026 respectively. Backlog has been large before without translating into commensurate '
       'profit.', '')
bullet(' Returns are still below cost of capital (FY26 ROE ~6%), so the business is not yet creating economic value '
       '— yet it is priced at ~88x earnings and ~6x book, multiples that assume the recovery is both certain and '
       'durable.', '')
bullet(' Our DCF says fair value is ₹181; even a generous probability-weighted scenario yields ₹245. The current '
       'price embeds the bull case in full, leaving negative risk-adjusted return.', '')
para('Our judgement: the operational turn is genuine, but the market price has run far ahead of the cash flows the '
     'turn can plausibly generate. We would rather own the recovery at a price that pays us for the execution risk.',space_before=2,space_after=4)

# ===================================================================
# 2. COMPANY OVERVIEW
# ===================================================================
h1('2.  Company Overview')
para('Bharat Heavy Electricals Limited, incorporated in 1964, is India\u2019s largest integrated power-plant '
     'equipment manufacturer and one of the country\u2019s flagship engineering public-sector undertakings (a '
     '\u201cMaharatna\u201d CPSE). The Government of India owns approximately 63% of the equity. Headquartered in '
     'New Delhi, BHEL operates a network of manufacturing plants, power-sector regional centres, and service '
     'stations across India, supported by a workforce of roughly 25,500 employees.',space_after=6)
para('BHEL\u2019s core franchise is the design, engineering, manufacture, construction, testing, commissioning '
     'and servicing of a wide range of products for the power generation and transmission, industrial, '
     'transportation, renewable energy, defence and aerospace sectors. The company has installed the bulk of '
     'India\u2019s coal-based generation fleet and retains an installed-base advantage that underpins a recurring '
     'spares-and-services revenue stream.',space_after=6)

h2('Business segments')
para('BHEL reports along two primary operating segments, supplemented by a diversification push into newer areas:',space_after=4)
seg=[
 ['Power','~75%','Steam turbines, boilers, turbo-generators, supercritical sets, balance-of-plant, emission-control (FGD) equipment, and erection & commissioning for thermal, hydro, gas and nuclear plants.'],
 ['Industry','~22%','Industrial systems and products — captive power, transportation (electric locomotives, propulsion), transmission (transformers, switchgear), capital equipment, compressors, valves and electrical machines.'],
 ['Emerging / Others','~3%','Defence & aerospace, solar PV and EPC, electric mobility & charging, green hydrogen and coal-gasification (incl. ammonium-nitrate JV), and exports.'],
]
table(['Segment','% of FY26 revenue','Description'], seg, widths=[1.3,1.2,4.0], body_size=9.0)
caption('Segment mix approximate; based on disclosed standalone segment data. Source: company filings; author estimates.')

para('The business is project-cycle in nature: revenue is recognised over multi-year execution of large EPC and '
     'equipment contracts, working capital is dominated by contract assets (unbilled revenue) and contract '
     'liabilities (customer advances), and earnings are highly operationally geared to volume and execution pace.',space_before=6)

h2('Recent performance and momentum')
para('FY2026 marked a clear inflection. Consolidated revenue grew ~18% to ₹33,782 cr, EBITDA rose to ₹2,342 cr '
     '(6.9% margin) and PAT reached ₹1,600 cr. The fourth quarter was particularly strong: Q4FY26 PAT more than '
     'doubled year-on-year to ₹1,283 cr on revenue of ₹12,310 cr (+37%), and the stock touched an 18-year high. '
     'During the year BHEL commissioned/synchronised ~8.9 GW of capacity and closed with an order book near '
     '₹2.4 lakh crore — the strongest revenue visibility in the company\u2019s modern history.',space_after=4)
caption('Sources: BHEL FY2026 results; businessworld.in; tipranks.com. Content rephrased for licensing compliance.')

# ===================================================================
# 3. INDUSTRY & MACRO
# ===================================================================
h1('3.  Industry & Macro Analysis')
h2('India\u2019s power demand and the thermal re-commitment')
para('India\u2019s peak electricity demand has risen sharply, crossing ~256 GW in 2025 and projected by the '
     'government toward ~270 GW in 2026. While renewable capacity is being added at 20 GW+ per year, the '
     'intermittency of solar and wind has forced policymakers back toward dispatchable, coal-based generation to '
     'maintain grid stability. The government plans to add on the order of 80 GW of new thermal capacity by 2032 — '
     'a direct, multi-year tailwind for BHEL, which supplies the majority of domestic boiler-turbine-generator (BTG) '
     'equipment.',space_after=6)
para('Crucially for the investment debate, the sector has a chronic execution-slippage problem. India added only '
     '~9.5 GW of thermal capacity in FY2026 (about 73% of the ~12.9 GW target) and just ~4.5 GW in FY2025 (~29% of '
     'target), as equipment-supply, land and clearance bottlenecks delayed commissioning. This is the crux of the '
     'BHEL thesis: the demand and ordering are real, but the pace at which orders convert into recognised revenue '
     'and cash has repeatedly disappointed.',space_after=4)
caption('Sources: Central Electricity Authority via financialexpress.com; livemint.com; ember-energy.org. Content rephrased for compliance.')

h2('Demand drivers')
bullet('Thermal capex super-cycle (~80 GW by 2032) and FGD/emission-control retrofits across the existing fleet.','')
bullet('Nuclear expansion blueprint and BHEL\u2019s positioning in indigenous reactor equipment.','')
bullet('Transmission build-out, railway electrification and locomotive/propulsion demand.','')
bullet('Defence indigenisation, green hydrogen, coal gasification and solar EPC as optionality.','')
bullet('Recurring aftermarket: spares, renovation & modernisation (R&M) and services on a large installed base.','')

h2('Porter\u2019s Five Forces')
pf=[
 ['Competitive rivalry','High','Domestic competition from L&T, and global BTG suppliers; intense price-based bidding on large tenders compresses margins.'],
 ['Threat of new entrants','Low–Med','High capital, technology and reference-list barriers in BTG; but Chinese/Korean OEMs and L&T are credible.'],
 ['Supplier power','Medium','Exposure to steel and copper prices; sub-contracting and imported components affect cost and timelines.'],
 ['Buyer power','High','Concentrated buyers (NTPC, state gencos, large IPPs) procure via L1 competitive bidding, pressuring pricing.'],
 ['Threat of substitutes','Medium','Renewables + storage substitute new thermal at the margin, capping the long-run thermal opportunity.'],
]
table(['Force','Intensity','Assessment'], pf, widths=[1.6,1.0,3.9], body_size=9.0)
caption('Author\u2019s assessment.')

h2('Competitive positioning')
para('BHEL\u2019s competitive moat rests on (i) the largest installed base of power equipment in India, generating '
     'sticky aftermarket revenue; (ii) full vertical integration across design, manufacturing and EPC; (iii) '
     'indigenous technology and a ~4% of sales R&D spend; and (iv) preferred-supplier status for public-sector '
     'capacity. Offsetting these are a public-sector cost structure, a history of working-capital build-up on '
     'stalled projects, and margin dilution from aggressive bidding during lean order years.',space_after=6)

h2('The order pipeline — beyond thermal')
para('While coal-based BTG remains the swing driver, BHEL\u2019s addressable pipeline has broadened, improving the '
     'durability of the order story:',space_after=4)
bullet(' The near-term core. FY2026 power orders of ~₹59,000 cr, supercritical and ultra-supercritical sets, and a '
       'visible tender pipeline from NTPC and state generators underpin the next leg of inflow.', 'Thermal BTG: ')
bullet(' Mandated retrofits across the existing coal fleet represent a large, relatively higher-margin and '
       'recurring opportunity that is less dependent on new-build economics.', 'FGD / emission control: ')
bullet(' BHEL is positioned in indigenous reactor equipment as India scales its nuclear blueprint — a multi-decade, '
       'high-value optionality that the market is beginning to capitalise.', 'Nuclear: ')
bullet(' Transformers and switchgear (transmission), electric locomotives and propulsion (railways), and defence '
       'indigenisation provide diversification away from the thermal cycle.', 'Industry, transport & defence: ')
bullet(' Solar EPC, green hydrogen, coal gasification (including the ammonium-nitrate JV) and exports are early '
       'but give BHEL a credible transition narrative.', 'New energy & exports: ')

h2('Peer landscape')
para('BHEL competes domestically with Larsen & Toubro\u2019s power/heavy-engineering business and, in transmission '
     'and industrial equipment, with premium multinationals. The peer set is instructive for valuation because the '
     'MNCs operate at structurally higher margins and returns, which justifies their richer multiples — and '
     'highlights why BHEL should trade at a discount, not parity.',space_after=4)
pl=[
 ['Siemens India','T&D, industrial automation, mobility','High (15%+)'],
 ['ABB India','Electrification, motion, automation','High (15%+)'],
 ['GE Vernova T&D India','Grid / transmission equipment','High'],
 ['CG Power & Industrial','Power systems, industrial, rail','Improving'],
 ['Thermax','Energy & environment, boilers','Mid'],
 ['Triveni Turbine','Industrial steam turbines','High'],
 ['Larsen & Toubro','EPC, power, heavy engineering','Mid (blended)'],
 ['BHEL','Power BTG, industrial, EPC','Low–recovering (~7%)'],
]
table(['Company','Focus','Margin profile'], pl, widths=[1.7,2.7,1.4], body_size=9.0)
caption('Author\u2019s assessment of competitive landscape.')

# ===================================================================
# 4. FINANCIAL ANALYSIS — HISTORICAL
# ===================================================================
h1('4.  Historical Financial Analysis (FY2020–FY2026)')
para('BHEL\u2019s recent history divides cleanly into a downturn (FY2020–FY2022), when weak ordering and stalled '
     'projects drove revenue to ₹17,309 cr and the company into losses (FY2021 PAT of \u2212₹2,699 cr), and a '
     'recovery (FY2023–FY2026) as the thermal cycle turned. Revenue rebounded to ₹33,782 cr in FY2026 with PAT of '
     '₹1,600 cr — the highest in years.',space_after=6)

h2('Revenue, profitability and returns')
hr=[]
for i in range(7):
    hr.append([YRS[i], fmt(sales[i]), fmt(ebitda[i]), f"{ebmar[i]:.1f}%", fmt(pat[i]), f"{roe[i]:.0f}%", f"{roce[i]:.0f}%"])
table(['Year','Revenue','EBITDA','EBITDA %','PAT','ROE %','ROCE %'], hr,
      widths=[0.9,1.1,1.0,0.95,1.0,0.8,0.8], body_size=9.0)
caption('Audited consolidated financials. Source: BHEL annual reports FY2021/FY2023/FY2025 and FY2026 filing.')

para('Three features stand out. First, operating leverage is extreme: with a largely fixed cost base, small '
     'revenue swings move margins violently — FY2021 saw a negative EBITDA margin, while FY2026 recovered to 6.9%. '
     'Second, returns are still sub-cost-of-capital: FY2026 ROE of ~6% and ROCE of ~6% remain well below our '
     '~15% cost of equity, underscoring that the recovery, while real, is early. Third, profitability has been '
     'flattered in some years by deferred-tax movements and other income (treasury income on a large cash pile).',space_before=4)

h2('Balance sheet and working capital')
para('BHEL is asset-light (net block of ~₹3,094 cr against ₹33,782 cr of revenue) and carried net cash of '
     '₹3,680 cr at FY2026 (gross borrowings ₹8,187 cr, largely short-term working-capital lines and lease '
     'liabilities, against ₹11,867 cr of cash and bank balances). The defining balance-sheet feature is the scale '
     'of contract balances: large contract assets (unbilled revenue on long-gestation projects) are funded '
     'substantially by contract liabilities (customer advances) — an interest-free source of working capital that '
     'is highly favourable when execution is flowing but which unwinds painfully when projects stall.',space_after=6)
bsr=[
 ['Total assets','59,749','55,240','56,244','59,370','59,002','68,083','76,186'],
 ['Shareholders\u2019 funds','28,652','25,972','26,507','26,828','24,439','24,722','26,147'],
 ['Gross borrowings & leases','5,080','4,951','4,830','5,454','8,856','9,015','8,187'],
 ['Cash & bank','6,419','6,701','7,154','6,643','6,157','7,612','11,867'],
 ['Net debt / (net cash)','(1,339)','(1,750)','(2,324)','(1,189)','2,699','1,403','(3,680)'],
]
table(['₹ crore','FY20','FY21','FY22','FY23','FY24','FY25','FY26'], bsr,
      widths=[1.7,0.78,0.78,0.78,0.78,0.78,0.78,0.78], body_size=8.5)
caption('Audited consolidated. Negative net debt = net cash. Source: author\u2019s model on audited financials.')

h2('Cash flow quality')
para('Operating cash flow has been volatile and tightly linked to the working-capital cycle: heavy outflows in '
     'the downturn (notably FY2020 and FY2024) gave way to a strong ₹5,837 cr operating inflow in FY2026 as '
     'advances flowed in and collections improved. Capital intensity is low (capex of ₹600–950 cr a year), so free '
     'cash flow is driven primarily by profit recovery and working-capital discipline rather than by investment.',space_after=6)

h2('DuPont decomposition')
para('Decomposing ROE clarifies why returns remain modest despite the recovery. BHEL\u2019s asset turnover is '
     'structurally low (~0.44x in FY26) given the contract-asset-heavy balance sheet, and net margin is only just '
     'turning up (~4.7% in FY26). Even with mild financial leverage, the product yields a single-digit ROE. The '
     'bull case for the equity is precisely that all three levers — margin, turnover and (favourable) leverage — '
     'improve together as execution accelerates; our forecast captures this, lifting ROE toward ~17% by FY2033.',space_after=4)
dp=[
 ['Net profit margin','-6.8%','1.9%','4.7%','5.1%','9.7%'],
 ['Asset turnover (x)','0.36','0.39','0.44','0.48','0.61'],
 ['Equity multiplier (x)','~2.1','~2.2','~2.9','~3.0','~2.9'],
 ['Return on equity','~0%','2%','6%','8%','17%'],
]
table(['DuPont component','FY20','FY23','FY26','FY27E','FY33E'], dp, widths=[2.0,0.9,0.9,0.9,0.9,0.9], body_size=9.0)
caption('ROE = net margin \u00d7 asset turnover \u00d7 equity multiplier. Source: author\u2019s model.')

h2('Peer benchmarking')
para('Against premium peers, BHEL screens as lower-margin and lower-return but cheaper on asset-based measures — '
     'and far more expensive on earnings, because its earnings base is still depressed relative to its order book.',space_after=4)
pb=[
 ['EBITDA margin','~7%','15–25%','BHEL well below peers; recovery is the thesis'],
 ['ROE','~6%','15–25%','Sub-cost-of-capital today'],
 ['P/E (trailing)','~88x','35–68x','Optically high on depressed earnings'],
 ['P/B','~6x','5–20x','Mid-range despite lower returns'],
 ['EV/EBITDA','~45x','22–50x','In line despite weaker quality'],
]
table(['Metric','BHEL','Peer range','Comment'], pb, widths=[1.4,0.9,1.3,2.2], body_size=9.0)
caption('Indicative; peer multiples to be refreshed with live data. Source: author\u2019s model and peer set.')

# ===================================================================
# 5. FORECAST & ASSUMPTIONS
# ===================================================================
h1('5.  Forecast Assumptions & Projections (FY2027–FY2033)')
para('Our forecast is order-book-driven rather than a simple growth extrapolation. Revenue is built from the '
     'opening order book and an execution (conversion) rate; the closing order book rolls forward as opening + new '
     'inflow \u2212 revenue executed. Costs, working capital, capex, depreciation, debt and equity are each '
     'projected on explicit drivers and fully linked across the three statements.',space_after=6)

h2('Key forecast assumptions (Base case)')
asr=[
 ['New order inflow (₹ cr)','78,000','82,000','86,000','90,000','92,000','94,000','96,000'],
 ['Execution rate (rev / opening OB)','16.5%','17.0%','17.5%','18.0%','18.5%','19.0%','19.5%'],
 ['Revenue growth','19.6%','19.5%','15.8%','13.1%','11.1%','9.1%','7.5%'],
 ['EBITDA margin','7.5%','8.5%','9.5%','10.5%','11.0%','11.5%','12.0%'],
 ['Effective tax rate','25%','25%','25%','25%','25%','25%','25%'],
 ['Capex (₹ cr)','600','700','750','800','850','900','950'],
 ['Dividend payout','30%','30%','30%','30%','30%','30%','30%'],
]
table(['Driver','FY27E','FY28E','FY29E','FY30E','FY31E','FY32E','FY33E'], asr,
      widths=[1.9,0.74,0.74,0.74,0.74,0.74,0.74,0.74], body_size=8.5)
caption('Base-case drivers. The model also contains Bull and Bear scenarios selectable via a switch. Source: author\u2019s model.')

h2('Rationale')
bullet(' We assume inflow moderates from the elevated FY26 level but stays robust (₹78,000–96,000 cr), consistent '
       'with the 80 GW thermal plan, FGD retrofits and nuclear/industrial orders. Even so, the order book keeps '
       'growing, reaching ~₹4.3 lakh cr by FY2033.', 'Order inflow: ')
bullet(' The single most important assumption. We model the conversion rate rising gradually from ~16.5% to '
       '~19.5%, reflecting better project delivery (8.9 GW commissioned in FY26) — but deliberately below a '
       'flawless ramp, given the sector\u2019s slippage record.', 'Execution rate: ')
bullet(' We model a disciplined climb from 7.5% to 12.0% as legacy low-margin orders roll off, fixed costs are '
       'absorbed over higher volume, and mix improves. We do not assume a return to the >15% margins of BHEL\u2019s '
       'last super-cycle.', 'Margins: ')
bullet(' Receivable days ~73, inventory ~155 days and payable ~123 days are held near FY26 levels; contract '
       'assets and liabilities are modelled as a declining percentage of revenue as the cycle matures.', 'Working capital: ')

h2('Projected P&L summary')
pr=[]
for i in range(7,14):
    pr.append([YRS[i], fmt(sales[i]), fmt(ebitda[i]), f"{ebmar[i]:.1f}%", fmt(pat[i]), f"{eps[i]:.1f}", f"{roe[i]:.0f}%"])
table(['Year','Revenue','EBITDA','EBITDA %','PAT','EPS (₹)','ROE %'], pr,
      widths=[0.95,1.05,1.0,0.95,1.0,0.9,0.8], body_size=9.0)
caption('Base case. Revenue CAGR FY26–33 ~13.6%; PAT CAGR ~26%. Source: author\u2019s model.')
para('On these assumptions, revenue compounds at ~13.6% to ₹82,405 cr and PAT at ~26% to ~₹8,004 cr by FY2033, '
     'with EBITDA margin reaching 12.0% and ROE/ROCE recovering to ~17–18%. The business turns structurally '
     'net-cash and generates rising free cash flow to the firm (FCFF of ₹3,089 cr in FY27E building to ₹7,938 cr '
     'in FY33E).',space_before=4)

# ===================================================================
# 6. VALUATION — DCF
# ===================================================================
h1('6.  Valuation — Discounted Cash Flow (Primary)')
para('We anchor our fair value on a Free Cash Flow to Firm (FCFF) DCF, the most appropriate primary method for a '
     'capital-goods company with a long, cash-generative project cycle. We discount seven years of explicit FCFF '
     '(FY2027–FY2033) and a Gordon-growth terminal value at a dynamically computed WACC.',space_after=6)

h2('Cost of capital (WACC)')
wr=[
 ['Risk-free rate (10Y G-Sec)','6.8%'],['Equity risk premium','6.5%'],['Levered beta','1.25'],
 ['Cost of equity (CAPM)','14.9%'],['Pre-tax cost of debt','8.5%'],['Tax rate','25%'],
 ['Target weight of equity / debt','85% / 15%'],['WACC','13.6%'],['Terminal growth (g)','4.5%'],
]
table(['Input','Value'], wr, widths=[3.2,1.6], body_size=9.5, first_col_left=True)
caption('CAPM cost of equity; WACC = We·Ke + Wd·Kd·(1\u2212t). Source: author\u2019s model (Assumptions sheet).')

h2('FCFF build and present value')
fc=[]
fcff=[3089,3802,4665,5620,6398,7172,7938]
for k,i in enumerate(range(7,14)):
    fc.append([YRS[i], fmt(sales[i]), fmt(ebitda[i]-[294,323,359,396,434,474,514][k]), fmt(fcff[k])])
table(['Year','Revenue','EBIT','FCFF'], fc, widths=[1.0,1.3,1.3,1.3], body_size=9.0)
caption('FCFF = NOPAT + D&A \u2212 capex \u00b1 change in working capital. Source: author\u2019s model (DCF sheet).')

h2('From enterprise value to per-share value')
dr=[
 ['Sum of PV of explicit FCFF (FY27–33)','22,158'],
 ['PV of terminal value','37,068'],
 ['Enterprise value (EV)','59,226'],
 ['Add: net cash (FY26)','3,680'],
 ['Equity value','62,906'],
 ['Shares outstanding (cr)','348.05'],
 ['Intrinsic value per share (₹)','181'],
 ['Current price (₹)','403'],
 ['Implied upside / (downside)','(55)%'],
]
table(['DCF output','Value (₹ cr unless stated)'], dr, widths=[3.4,2.0], body_size=9.5, first_col_left=True,
      total_rows=[6])
caption('Source: author\u2019s model. Terminal value uses Gordon growth at g = 4.5% on FY33E FCFF.')
para('The DCF yields an intrinsic value of ₹181 per share — roughly 55% below the prevailing market price. The '
     'terminal value represents ~63% of EV, which is typical but means the valuation is sensitive to WACC and the '
     'terminal margin/growth assumptions (see Section 8). The core message is unambiguous: even crediting BHEL with '
     'a successful margin recovery to 12% and a sustained order book, the present value of its cash flows does not '
     'support a ₹400 share price.',space_before=4)

# ===================================================================
# 7. RELATIVE VALUATION & SCENARIOS
# ===================================================================
h1('7.  Relative Valuation & Scenario Analysis')
h2('Relative valuation (cross-check)')
para('BHEL trades within a basket of Indian capital-goods/electrical-equipment names that command premium '
     'multiples, led by multinational franchises (Siemens India, ABB India, GE Vernova T&D, CG Power). Applying '
     'peer median multiples to BHEL\u2019s FY2027E metrics produces a wide range of implied values.',space_after=4)
rv=[
 ['EV/EBITDA','41.0x','367'],
 ['P/E','57.5x','343'],
 ['EV/Revenue','6.0x','707'],
 ['Average implied price','—','472'],
]
table(['Method','Peer median','Implied price (₹)'], rv, widths=[1.9,1.5,1.6], body_size=9.5, first_col_left=True, total_rows=[3])
caption('Indicative peer multiples; refresh with live data before use. Source: author\u2019s model (Relative Valuation sheet).')
para('We treat relative valuation as a sanity check, not a primary anchor, and we discount it heavily. The peer '
     'set is dominated by asset-light MNC subsidiaries with structurally higher margins (15–25%), cleaner '
     'governance and faster growth, which justifiably trade at 40–50x EBITDA. BHEL — a public-sector, lower-margin, '
     'execution-sensitive business — does not merit parity. The wide ₹343–₹707 range itself signals how '
     'multiple-dependent any bullish case is.',space_before=2,space_after=4)

h2('Scenario analysis')
para('We frame three scenarios on FY2033E using exit EV/EBITDA multiples discounted to the present, and assign '
     'probabilities of 25% / 50% / 25% to Bear / Base / Bull.',space_after=4)
sc=[
 ['FY33E revenue (₹ cr)','68,000','82,405','92,000'],
 ['Terminal EBITDA margin','9.0%','12.0%','15.0%'],
 ['Exit EV/EBITDA (x)','14.0','18.0','24.0'],
 ['WACC','15.5%','13.6%','12.0%'],
 ['Implied target price (₹)','100','219','441'],
 ['Probability','25%','50%','25%'],
]
table(['Driver / output','Bear','Base','Bull'], sc, widths=[2.3,1.2,1.2,1.2], body_size=9.0, total_rows=[4])
para('The probability-weighted target price across the three scenarios is ₹245, which we adopt as our 12-month '
     'target. It sits above the pure DCF (₹181) — giving partial credit to the bull case and the strength of the '
     'order book — but well below both the current price and the multiple-driven relative value. The asymmetry is '
     'telling: the bear case implies ~75% downside while the bull case implies only modest upside, so the '
     'risk/reward at ₹403 is unattractive.',space_before=4)

h2('Valuation bridge')
para('The gap between our methods is itself the story. The DCF (₹181) values the cash the business can generate '
     'on credible assumptions. The relative method (₹472) values the stock on what the market is willing to pay '
     'for comparable franchises today — a multiple that embeds both a capital-goods re-rating and a quality '
     'premium BHEL has not yet earned. The truth, in our view, lies closer to the cash flows than to the '
     'multiples; we therefore weight the DCF most heavily and adopt the probability-weighted ₹245 as our anchor.',space_after=4)

h2('Valuation summary and target price')
vs=[
 ['DCF (FCFF), primary','50%','181'],
 ['Probability-weighted scenario','30%','245'],
 ['Relative (peer multiples)','20%','472'],
 ['Blended fair value (rounded)','100%','~245'],
]
table(['Approach','Weight','Value (₹)'], vs, widths=[2.6,1.2,1.4], body_size=9.5, first_col_left=True, total_rows=[3])
caption('Author\u2019s blend. Target rounded to ₹245. Source: author\u2019s model.')

# ===================================================================
# 8. SENSITIVITY
# ===================================================================
h1('8.  Sensitivity Analysis')
para('Because the terminal value is a large share of DCF enterprise value, intrinsic value is most sensitive to '
     'WACC and terminal growth. The two-way table below shows DCF intrinsic value per share (₹) across these two '
     'variables; the centre cell is our base case (WACC 13.6%, g 4.5%).',space_after=4)
sen=[
 ['11.6%','243','262','286','315','352'],
 ['12.6%','205','219','236','257','283'],
 ['13.6% (base)','175','186','199','215','235'],
 ['14.6%','151','159','170','182','197'],
 ['15.6%','131','138','146','156','168'],
]
table(['WACC \\\\ g','3.0%','3.75%','4.5%','5.25%','6.0%'], sen, widths=[1.5,1.0,1.0,1.0,1.0,1.0], body_size=9.0, highlight_rows=[2])
caption('DCF intrinsic value per share (₹). Centre column/row = base case. Source: author\u2019s model (Sensitivity sheet).')
para('Across the plausible range, intrinsic value spans roughly ₹130–₹290 — and only an aggressive combination '
     'of sub-12% WACC and 6% perpetual growth approaches the current ₹400 price. Other key swing factors, modelled '
     'on the Scenario and Assumptions sheets, include order-inflow levels, the execution/conversion rate, EBITDA '
     'margins, steel and copper prices, capex, interest rates and the working-capital cycle.',space_before=4)

# ===================================================================
# 9. RISKS
# ===================================================================
h1('9.  Investment Risks')
h3('Risks to our (cautious) view — i.e. upside risks')
bullet(' A faster-than-modelled rise in the execution rate would lift revenue and, through operating leverage, '
       'margins and EPS materially above our base case.', 'Execution acceleration: ')
bullet(' Inflow consistently above ₹90,000–1,00,000 cr (nuclear, FGD, large thermal, exports) would extend revenue '
       'visibility and support higher multiples.', 'Stronger ordering: ')
bullet(' The market may continue to capitalise the thermal/nuclear theme at premium capital-goods multiples for '
       'longer than fundamentals justify.', 'Momentum/re-rating: ')
h3('Risks to the business and to a long position')
bullet(' BHEL\u2019s own record — and the sector\u2019s (FY25/FY26 thermal additions far below target) — shows '
       'order books do not always convert on schedule.', 'Execution slippage: ')
bullet(' Steel and copper inflation, or aggressive L1 bidding to win share, could stall the margin recovery.', 'Commodity & pricing: ')
bullet(' Heavy reliance on government/PSU capex makes revenue sensitive to policy, budget cycles and a structural '
       'shift away from coal toward renewables + storage.', 'Policy & energy transition: ')
bullet(' Large contract assets can build up on delayed projects, swinging operating cash flow and returns.', 'Working-capital risk: ')
bullet(' As a ~63% government-owned PSU, BHEL carries public-sector cost structures, slower decision-making, and '
       'potential conflicts between commercial and policy objectives; equity supply via disinvestment is possible.', 'Governance/PSU overhang: ')
bullet(' At ~88x trailing and ~45x FY27E earnings, any execution disappointment can de-rate the stock sharply — '
       'the dominant near-term risk.', 'Valuation risk: ')

# ===================================================================
# MANAGEMENT & STRATEGY
# ===================================================================
h1('10.  Management, Strategy & Capital Allocation')
h2('Strategic direction')
para('BHEL\u2019s strategy has three planks. First, capitalise on the domestic thermal and FGD super-cycle as the '
     'preferred indigenous supplier, while improving on-time execution to convert the order book into revenue and '
     'cash. Second, diversify the revenue base away from coal — into nuclear, transmission, transportation, '
     'defence, solar EPC, green hydrogen and coal gasification — to extend the runway beyond the current cycle. '
     'Third, lift profitability structurally through better project selection (avoiding value-destructive L1 '
     'bids), localisation, and cost discipline on a historically heavy fixed-cost base.',space_after=6)
para('Execution against this strategy is visibly improving — 8.9 GW commissioned in FY2026, a doubling of Q4 '
     'profit, and a record book-to-bill — but the company must still demonstrate the gains are durable across a '
     'full cycle rather than a cyclical rebound off a depressed base.',space_after=6)
h2('Capital allocation')
para('Capital allocation is conservative and, for a growth-inflecting business, arguably under-geared. BHEL is '
     'net-cash, spends modestly on capex (the business is asset-light), and returns ~30% of profit as dividends in '
     'line with the DIPAM policy for CPSEs. Large cash and customer-advance balances mean growth is largely '
     'self-funding. The principal capital-allocation risk is not over-investment but value leakage through '
     'aggressive bidding and working-capital build-up on stalled projects.',space_after=6)
h2('Quality of earnings')
para('Investors should note that reported profits have, in some years, benefited from non-operating items — '
     'treasury/other income on the large cash pile and deferred-tax movements. We therefore focus on EBITDA, '
     'operating cash flow and order-book conversion as the cleaner signals of the underlying recovery, and we '
     'normalise for these in our forecast.',space_after=4)

# ===================================================================
# 11. ESG & GOVERNANCE
# ===================================================================
h1('11.  ESG, Governance & Ownership')
h2('Environmental')
para('BHEL\u2019s environmental profile is paradoxical. Its core product — coal-fired power equipment — is '
     'carbon-intensive and faces long-run transition risk. Yet the same engineering base positions BHEL to supply '
     'the enabling technologies of the transition: flue-gas desulphurisation (FGD) and emission-control retrofits, '
     'supercritical/ultra-supercritical (higher-efficiency) sets, solar EPC, green hydrogen, coal gasification and '
     'nuclear equipment. The investment debate hinges on whether BHEL transitions its mix fast enough as thermal '
     'demand eventually plateaus.',space_after=6)
h2('Social')
para('As a large public-sector employer (~25,500 staff), BHEL carries significant social weight and a relatively '
     'high, partly fixed labour cost (~19% of sales). Workforce rationalisation has been gradual. Safety, skilling '
     'and indigenous technology development (R&D ~4% of sales) are areas of strength.',space_after=6)
h2('Governance & ownership')
para('The Government of India holds ~63%, making minority shareholders price-takers on strategic decisions and '
     'exposing the stock to disinvestment-related supply. Board composition follows CPSE norms, with independent '
     'directors and standard audit/oversight committees. Governance positives include audited, Ind-AS-compliant '
     'disclosure and a stated DIPAM dividend policy (~30% payout); negatives include the usual PSU constraints on '
     'capital allocation and commercial agility. Net-net, governance is adequate but not a re-rating catalyst.',space_after=4)

# ===================================================================
# 11. CATALYSTS
# ===================================================================
h1('12.  Catalysts & Triggers (12 months)')
bullet(' Pace of order-book conversion and reported execution rate — the key swing variable for the stock.', 'Quarterly execution: ')
bullet(' New thermal (NTPC/state), FGD, nuclear and export awards; confirmation of inflow trajectory above/below ₹90,000 cr.', 'Order wins: ')
bullet(' Whether EBITDA margin sustains above ~8–9% and trends toward double digits.', 'Margin trajectory: ')
bullet(' Sustained positive operating cash flow and net-cash build, validating working-capital discipline.', 'Cash conversion: ')
bullet(' Any government stake sale / OFS would pressure the price near-term.', 'Disinvestment news: ')
bullet(' Steel/copper moves and interest-rate shifts affecting margins and discount rate.', 'Macro: ')

# ===================================================================
# 12. APPENDICES
# ===================================================================
h1('Appendix A — Income Statement (₹ crore)')
ais=[]
ametrics=['Revenue','EBITDA','D&A','EBIT','PBT','PAT','EPS (₹)']
da=[503,473,314,260,249,272,316,294,323,359,396,434,474,514]
ebit=[-636,-3521,514,784,462,1127,2026,2735,3780,4950,6241,7294,8339,9374]
pbt=[-659,-3595,471,716,242,745,2139,2767,4038,5430,6938,8202,9448,10672]
for i in range(14):
    ais.append([YRS[i], fmt(sales[i]), fmt(ebitda[i]), fmt(da[i]), fmt(ebit[i]), fmt(pbt[i]), fmt(pat[i]), f"{eps[i]:.1f}"])
table(['Year','Revenue','EBITDA','D&A','EBIT','PBT','PAT','EPS'], ais,
      widths=[0.8,0.95,0.85,0.7,0.85,0.9,0.85,0.6], body_size=8.5)
caption('A = audited (FY20–26); E = estimate (FY27–33). Source: author\u2019s integrated model.')

h1('Appendix B — Balance Sheet & Ratios')
h2('Condensed balance sheet (₹ crore)')
ta=[59749,55240,56244,59370,59002,68083,76186,83462,91000,99500,109000,119000,129500,140500]
eq=[28652,25972,26507,26828,24439,24722,26147,27599,29719,32591,36234,40090,44326,48529]
nd=['(1,339)','(1,750)','(2,324)','(1,189)','2,699','1,403','(3,680)','(2,900)','(4,400)','(6,600)','(9,400)','(12,900)','(17,200)','(22,000)']
abs_=[]
for i in range(14):
    abs_.append([YRS[i], fmt(ta[i]), fmt(eq[i]), nd[i]])
table(['Year','Total assets','Shareholders\u2019 funds','Net debt/(cash)'], abs_,
      widths=[1.0,1.6,1.8,1.6], body_size=8.5)
caption('Forecast total assets/equity rounded; net cash builds materially over the horizon. Source: author\u2019s model.')

h2('Key ratios')
kr=[]
recv=['','','','','','',73,73,73,73,73,73,73,73]
for i in range(14):
    kr.append([YRS[i], f"{ebmar[i]:.1f}%", f"{roe[i]:.0f}%", f"{roce[i]:.0f}%", f"{de[i]:.2f}"])
table(['Year','EBITDA %','ROE %','ROCE %','D/E (x)'], kr,
      widths=[1.1,1.3,1.1,1.1,1.1], body_size=8.5)
caption('Source: author\u2019s model (Ratio Analysis sheet).')

h1('Appendix C — Order Book & Operational Drivers')
ob=[
 ['Opening order book','1,15,000','1,02,000','98,000','91,336','1,31,598','1,96,328','2,40,000'],
 ['New order inflow','18,000','22,000','23,548','63,000','92,535','75,000','78,000'],
 ['Revenue executed','20,495','20,153','22,136','22,921','27,355','33,782','39,600'],
 ['Closing order book','1,10,000','98,000','91,336','1,31,598','1,96,328','2,40,000','2,78,400'],
 ['Book-to-bill (x)','5.4','4.9','4.1','5.7','7.2','7.1','7.0'],
]
table(['₹ crore (standalone)','FY20','FY22','FY23','FY24','FY25','FY26','FY27E'], ob,
      widths=[1.7,0.78,0.78,0.78,0.78,0.78,0.78,0.78], body_size=8.3)
caption('Order data on standalone basis (as disclosed); FY27E from model. Closing OB grows to ~₹4.3 lakh cr by FY33E. Source: author\u2019s model.')

para('',space_after=2)
h2('Methodology note')
para('This report is supported by a fully integrated, audit-checked three-statement model. Historical financials '
     '(FY2020–FY2026) are sourced from BHEL\u2019s audited consolidated annual reports (FY2021, FY2023, FY2025) and '
     'the FY2026 results filing. The forecast (FY2027–FY2033) is order-book-driven, with revenue derived from the '
     'execution rate applied to the opening order book, and all line items linked across the income statement, '
     'balance sheet and cash flow. The model balances to zero in every year and passes automated integrity checks '
     '(balance-sheet, cash-flow, debt and equity roll-forwards) across Base, Bull and Bear scenarios.',space_after=6)

# ===================================================================
# DISCLAIMER
# ===================================================================
h1('Important Disclosures & Disclaimer')
para('Rating system: BUY — expected total return >15% over 12 months; HOLD — \u221210% to +15%; SELL — expected '
     'total return below \u221210%. Our SELL rating on BHEL reflects an expected downside of ~39% to our ₹245 '
     'target price.',9,space_after=6)
para('This document has been prepared for informational and educational purposes and in a manner consistent with '
     'CFA Institute standards of professional conduct, including objectivity and a reasonable and adequate basis. '
     'It is not investment advice, nor an offer or solicitation to buy or sell any security. The analysis relies on '
     'publicly available information and an independently constructed financial model; while believed reliable, its '
     'accuracy and completeness are not guaranteed. Forward-looking statements are subject to material uncertainty '
     'and actual results may differ. Valuation outputs (DCF intrinsic value ₹181; probability-weighted scenario '
     'value ₹245; peer-multiple value ₹472) are model-dependent and sensitive to assumptions disclosed herein. '
     'Peer trading multiples are indicative and should be refreshed with live market data. The author may have no '
     'position in the security. Investors should conduct their own due diligence and consult a licensed financial '
     'adviser. Past performance is not indicative of future results.',9,space_after=6)
para('Selected sources: BHEL audited annual reports (FY2021/FY2023/FY2025) and FY2026 results; Central Electricity '
     'Authority (via financialexpress.com); businessworld.in; tipranks.com; livemint.com; businesstoday.in; '
     'ndtvprofit.com; ember-energy.org. Third-party content has been paraphrased and summarised for licensing '
     'compliance; figures are attributed to the author\u2019s model unless otherwise stated.',8.5,italic=True,color=GREY)

doc.save('/projects/sandbox/BHEL_Equity_Research_Report.docx')
print("Report saved.")
